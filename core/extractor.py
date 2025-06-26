import os
import fitz
from PIL import Image
import io
from typing import Callable, Dict, Any, Optional, List, Tuple
from abc import ABC, abstractmethod
import google.generativeai as genai
from openai import OpenAI
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument
from core.config import GeminiConfig, OpenAIConfig

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"
TEMP_IMAGE_DIR = os.path.join(OUTPUT_DIR, "temp_images")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(TEMP_IMAGE_DIR, exist_ok=True)

class ExtractorStrategy(ABC):
    @abstractmethod
    def extract_text(self, file_path: str, progress_callback: Callable[[str], None]) -> str:
        pass

class LangchainPDFExtractor(ExtractorStrategy):
    def extract_text(self, file_path: str, progress_callback: Callable[[str], None]) -> str:
        progress_callback("Starting extraction with Custom algorithm...")
        progress_callback("Loading PDF document...")
        
        try:
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            
            progress_callback(f"PDF loaded into {len(documents)} document sections.")
            
            extracted_text = ""
            for i, doc_chunk in enumerate(documents):
                extracted_text += doc_chunk.page_content + "\n\n"
                progress_callback(f"Extracted text from section {i+1}/{len(documents)}.")
            
            progress_callback("Extraction complete.")
            return extracted_text
        except Exception as e:
            error_message = f"Error loading PDF: {str(e)}"
            progress_callback(error_message)
            return f"[Error: {error_message}]"

class GeminiPDFExtractor(ExtractorStrategy):
    def __init__(self, model_name: str):
        self.model_name = model_name
        self.config = GeminiConfig()
    
    def extract_text(self, file_path: str, progress_callback: Callable[[str], None]) -> str:
        progress_callback(f"Starting extraction with Gemini AI using {self.model_name}...")
        
        if not self.config.is_available():
            return "[Error: Gemini API key not configured]"
        
        self.config.configure()
        
        pdf_document = None
        try:
            progress_callback("Opening PDF document...")
            pdf_document = fitz.open(file_path)
        except Exception as e:
            progress_callback(f"Error opening PDF: {e}")
            return f"[Error: {str(e)}]"

        total_pages = len(pdf_document)
        progress_callback(f"PDF has {total_pages} pages.")
        
        all_pages_text_content = []
        model = genai.GenerativeModel(self.model_name)

        for page_num in range(total_pages):
            current_page_info = f"Page {page_num + 1}/{total_pages}"
            progress_callback(f"Processing {current_page_info}...")
            
            try:
                page = pdf_document.load_page(page_num)
                
                progress_callback(f"{current_page_info}: Converting to high-resolution image...")
                pix = page.get_pixmap(dpi=300)
                img_bytes = pix.tobytes("png")
                
                progress_callback(f"{current_page_info}: Sending image to Gemini...")
                
                prompt = "Extract all text content from this image. Preserve formatting like paragraphs and line breaks where possible."
                image_part = {"mime_type": "image/png", "data": img_bytes}
                
                response = model.generate_content([prompt, image_part])
                page_text = response.text
                
                if page_text.strip():
                    all_pages_text_content.append(page_text)
                else:
                    all_pages_text_content.append(f"[No text found on {current_page_info}]")
                    
                progress_callback(f"{current_page_info}: Text extraction complete.")

            except Exception as e:
                error_message = f"Error processing {current_page_info}: {str(e)}"
                progress_callback(error_message)
                all_pages_text_content.append(f"[{error_message}]")
                
        if pdf_document:
            pdf_document.close()
            
        progress_callback("Combining text from all pages...")
        final_text = "\n\n--- Page Break ---\n\n".join(all_pages_text_content)
        progress_callback("Extraction complete.")
        return final_text

class OpenAIPDFExtractor(ExtractorStrategy):
    def __init__(self, model_name: str):
        self.model_name = model_name
        self.config = OpenAIConfig()
    
    def extract_text(self, file_path: str, progress_callback: Callable[[str], None]) -> str:
        progress_callback(f"Starting extraction with OpenAI using {self.model_name}...")
        
        if not self.config.is_available():
            return "[Error: OpenAI API key not configured]"
        
        self.config.configure()
        client = OpenAI(api_key=self.config.api_key)
        
        pdf_document = None
        try:
            progress_callback("Opening PDF document...")
            pdf_document = fitz.open(file_path)
        except Exception as e:
            progress_callback(f"Error opening PDF: {e}")
            return f"[Error: {str(e)}]"

        total_pages = len(pdf_document)
        progress_callback(f"PDF has {total_pages} pages.")
        
        all_pages_text_content = []

        for page_num in range(total_pages):
            current_page_info = f"Page {page_num + 1}/{total_pages}"
            progress_callback(f"Processing {current_page_info}...")
            
            try:
                page = pdf_document.load_page(page_num)
                
                progress_callback(f"{current_page_info}: Converting to high-resolution image...")
                pix = page.get_pixmap(dpi=300)
                img_bytes = pix.tobytes("png")
                
                progress_callback(f"{current_page_info}: Sending image to OpenAI...")
                
                with open(os.path.join(TEMP_IMAGE_DIR, f"temp_page_{page_num}.png"), "wb") as f:
                    f.write(img_bytes)
                
                with open(os.path.join(TEMP_IMAGE_DIR, f"temp_page_{page_num}.png"), "rb") as img_file:
                    response = client.chat.completions.create(
                        model=self.model_name,
                        messages=[
                            {"role": "system", "content": "Extract all text content from this image. Preserve formatting like paragraphs and line breaks where possible."},
                            {"role": "user", "content": [
                                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64.b64encode(img_file.read()).decode('utf-8')}"}}
                            ]}
                        ]
                    )
                
                page_text = response.choices[0].message.content
                
                if page_text.strip():
                    all_pages_text_content.append(page_text)
                else:
                    all_pages_text_content.append(f"[No text found on {current_page_info}]")
                    
                progress_callback(f"{current_page_info}: Text extraction complete.")

            except Exception as e:
                error_message = f"Error processing {current_page_info}: {str(e)}"
                progress_callback(error_message)
                all_pages_text_content.append(f"[{error_message}]")
                
            # Remove temp file
            try:
                os.remove(os.path.join(TEMP_IMAGE_DIR, f"temp_page_{page_num}.png"))
            except:
                pass
                
        if pdf_document:
            pdf_document.close()
            
        progress_callback("Combining text from all pages...")
        final_text = "\n\n--- Page Break ---\n\n".join(all_pages_text_content)
        progress_callback("Extraction complete.")
        return final_text

class DocxExtractor(ExtractorStrategy):
    def extract_text(self, file_path: str, progress_callback: Callable[[str], None]) -> str:
        progress_callback("Starting DOCX text extraction...")
        
        try:
            progress_callback(f"Opening DOCX file: {os.path.basename(file_path)}...")
            document = DocxDocument(file_path)
            progress_callback("DOCX file opened. Extracting paragraphs...")
            
            full_text = []
            for i, para in enumerate(document.paragraphs):
                full_text.append(para.text)
                if (i + 1) % 20 == 0:
                    progress_callback(f"Processed {i+1} paragraphs...")
            
            progress_callback(f"Total {len(document.paragraphs)} paragraphs processed.")
            progress_callback("DOCX extraction complete.")
            return "\n".join(full_text)
            
        except Exception as e:
            error_message = f"Error reading DOCX file: {str(e)}"
            progress_callback(error_message)
            return f"[Error: {error_message}]"

class DocumentProcessor:
    def __init__(self):
        self.strategy = None
    
    def set_strategy(self, strategy: ExtractorStrategy) -> None:
        self.strategy = strategy
    
    def process(self, file_path: str, progress_callback: Callable[[str], None]) -> str:
        if self.strategy is None:
            return "[Error: No extraction strategy selected]"
        
        return self.strategy.extract_text(file_path, progress_callback)

class ExtractorFactory:
    @staticmethod
    def create_extractor(extractor_type: str, model_name: str = None) -> ExtractorStrategy:
        if extractor_type == "langchain":
            return LangchainPDFExtractor()
        elif extractor_type == "gemini":
            if not model_name:
                model_name = "gemini-1.5-flash-latest"
            return GeminiPDFExtractor(model_name)
        elif extractor_type == "openai":
            if not model_name:
                model_name = "gpt-4o"
            return OpenAIPDFExtractor(model_name)
        elif extractor_type == "docx":
            return DocxExtractor()
        else:
            raise ValueError(f"Unknown extractor type: {extractor_type}")

def get_pdf_first_page_preview(pdf_path: str) -> Optional[bytes]:
    try:
        doc = fitz.open(pdf_path)
        if len(doc) > 0:
            page = doc.load_page(0)
            pix = page.get_pixmap(dpi=96)
            img_bytes = pix.tobytes("png")
            doc.close()
            return img_bytes
        doc.close()
    except Exception as e:
        print(f"Error generating PDF preview: {e}")
    return None

def cleanup_temporary_files(uploaded_file_path: Optional[str] = None) -> None:
    if uploaded_file_path and os.path.exists(uploaded_file_path):
        try:
            os.remove(uploaded_file_path)
        except Exception as e:
            print(f"Error removing uploaded file {uploaded_file_path}: {e}")

    for item_name in os.listdir(TEMP_IMAGE_DIR):
        item_path = os.path.join(TEMP_IMAGE_DIR, item_name)
        try:
            if os.path.isfile(item_path):
                os.remove(item_path)
        except Exception as e:
            print(f"Error removing temp image {item_path}: {e}")

import base64